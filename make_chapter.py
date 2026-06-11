"""Builds the book chapter as a Word document from the project's real results.

    python make_chapter.py

Reads tables from reports/ and figures from reports/png/, and writes
"Diabetes_Prediction_Book_Chapter.docx" to the project root.
"""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
REPORTS = ROOT / "reports"
PNG = REPORTS / "png"
OUT = ROOT / "Diabetes_Prediction_Book_Chapter.docx"


def metrics():
    return json.loads((REPORTS / "metrics_summary.json").read_text())


# --------------------------------------------------------------------------
# Document helpers
# --------------------------------------------------------------------------
def add_body(doc, text, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    return p


def add_figure(doc, png_name, caption, width=5.4):
    path = PNG / png_name
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.italic = True


def add_table_from_csv(doc, csv_name, caption, col_format=None, max_rows=None):
    df = pd.read_csv(REPORTS / csv_name)
    if max_rows:
        df = df.head(max_rows)
    if col_format:
        for col, fmt in col_format.items():
            if col in df.columns:
                df[col] = df[col].map(lambda v: fmt.format(v) if pd.notna(v) else "—")
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = str(c)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
            for p in cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.italic = True


# --------------------------------------------------------------------------
# Chapter content
# --------------------------------------------------------------------------
def build():
    M = metrics()
    ext = M["external_validation"]
    doc = Document()

    # base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Putting a Risk Score in Its Place: Geographic Context for a "
                      "Machine-Learning Model of Diabetes")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Rishil Shah")
    r.font.size = Pt(12)
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = aff.add_run("Department of Computer Science\n[Affiliation to be completed]\n"
                    "rishil13123@gmail.com")
    r.font.size = Pt(10)
    r.italic = True
    doc.add_paragraph()

    # ---- Abstract ----
    h = doc.add_paragraph(); r = h.add_run("Abstract"); r.bold = True; r.font.size = Pt(12)
    add_body(doc,
        "Most tutorials on diabetes prediction end the moment a classifier clears some "
        "accuracy threshold on the Pima Indians dataset. That is a strange place to stop, "
        "because the harder questions begin right after. Does the model hold up on people it "
        "was never trained on? Are its probabilities trustworthy enough to act on, or only to "
        "rank by? And what does a single number mean for a patient who lives somewhere the "
        "training data never described? This chapter works through one model end to end with "
        "those questions in mind. We build a calibrated tree-ensemble on the Pima data, but we "
        "spend most of our effort on what comes after the fit: an ablation that shows which "
        "design choices actually earned their keep, confidence intervals on every reported "
        "number, a significance test that quietly demolishes the case for a fancy ensemble, an "
        "external validation on a completely separate national survey, and SHAP attributions "
        "that we checked against clinical intuition rather than the other way round. The piece "
        "we think is genuinely new is geographic: we link each prediction to county-level "
        "prevalence from the CDC PLACES program and use a simple odds correction to re-express "
        "an individual score against the place that person lives. The internal model reaches a "
        "cross-validated ROC-AUC near 0.84; more interestingly, its discrimination survives "
        "almost intact when moved to NHANES, even as the base rate falls by more than half. We "
        "argue that this transportability is exactly what licenses the geographic adjustment, "
        "and we are candid about where the whole approach should not be trusted.")

    kw = doc.add_paragraph()
    r = kw.add_run("Keywords  "); r.bold = True; r.font.size = Pt(10)
    r = kw.add_run("diabetes risk prediction · model calibration · external validation · "
                   "SHAP · GeoAI · CDC PLACES · health geography · responsible machine learning")
    r.font.size = Pt(10); r.italic = True

    # ---- 1 Introduction ----
    doc.add_heading("1  Introduction", level=1)
    add_body(doc,
        "The Pima Indians Diabetes dataset has been a teaching staple for more than thirty "
        "years. Eight measurements, one binary outcome, 768 rows: it is small enough to fit in "
        "memory a thousand times over and clean enough that a first-year student can get a "
        "respectable number out of it in an afternoon. That convenience is also its trap. The "
        "dataset is so well-behaved as a benchmark that it invites people to treat prediction as "
        "the finish line, when in any real screening setting it is barely the starting one.")
    add_body(doc,
        "We came at this project from the other direction. The model was never going to be the "
        "interesting contribution — the literature already contains dozens of papers reporting "
        "areas under the curve in the high 0.70s to mid-0.80s on this exact data, and we did not "
        "expect to beat them by much, nor did we. What we wanted to know was whether a model "
        "trained on a few hundred women from one community in Arizona could say anything "
        "defensible about a person somewhere else entirely, and whether we could attach an honest "
        "measure of place to its output. The second question is where the geographic component "
        "comes in, and it is the part of this chapter we would point a reviewer to first.")
    add_body(doc,
        "The structure follows the order in which we actually worked. Section 2 sketches the "
        "background and names the prior work we leaned on. Section 3 describes the three data "
        "sources — Pima for training, NHANES for external validation, and CDC PLACES for "
        "geography. Section 4 lays out the methods, including the prior-correction formula that "
        "underpins the geographic claim. Section 5 reports results, and we have tried to report "
        "the unflattering ones with the same prominence as the flattering ones. Section 6 "
        "describes the interactive system we built around the model. Section 7 deals with ethics "
        "and limitations, which on this particular dataset are not an afterthought. Section 8 "
        "closes.")

    # ---- 2 Background ----
    doc.add_heading("2  Background and Related Work", level=1)
    add_body(doc,
        "The dataset originates with a 1988 study by Smith and colleagues, who used an early "
        "neural network they called ADAP to forecast the onset of diabetes in a population of "
        "adult women of Pima (Akimel O'odham) heritage. The features — number of pregnancies, "
        "plasma glucose, diastolic blood pressure, triceps skinfold thickness, serum insulin, "
        "body-mass index, a pedigree function summarising family history, and age — have been "
        "copied into teaching materials ever since. We say more about the ethics of this "
        "provenance in Section 7; for now it is enough to note that the data describe a specific "
        "community with a specific and painful relationship to diabetes research, and that this "
        "matters for how far any conclusion should travel.")
    add_body(doc,
        "On the modelling side we relied on well-worn tools rather than anything exotic. The "
        "ensemble of extremely randomised trees follows Geurts and colleagues; gradient boosting "
        "is represented by both the scikit-learn implementation and the XGBoost and LightGBM "
        "libraries. For turning raw scores into trustworthy probabilities we used Platt's "
        "sigmoid scaling. To compare two models without fooling ourselves we used the DeLong "
        "test for correlated ROC curves, and to explain individual predictions we used the "
        "SHAP framework of Lundberg and Lee. None of this is novel, and that is deliberate: the "
        "novelty we are claiming is in the geographic layer, not the estimator.")
    add_body(doc,
        "The geographic data come from PLACES, a collaboration between the CDC, the Robert Wood "
        "Johnson Foundation, and the CDC Foundation that publishes small-area estimates of "
        "chronic-disease measures for every county, place, and census tract in the United "
        "States. To our knowledge, coupling a personal Pima-trained risk model to PLACES "
        "prevalence through an explicit prior correction has not been done before, and it is the "
        "one thing in this chapter we would defend as genuinely new.")

    # ---- 3 Data ----
    doc.add_heading("3  Data", level=1)
    doc.add_heading("3.1  Training cohort: Pima", level=2)
    add_body(doc,
        "The training data are the standard 768 records, of which 268 (34.9 percent) carry a "
        "positive diabetes label. A quirk that catches many newcomers is that several columns "
        "use zero as a stand-in for a missing measurement. A body-mass index of zero or a "
        "diastolic blood pressure of zero is not a low reading; it is no reading at all. We treat "
        "zeros in glucose, blood pressure, skinfold thickness, insulin, and BMI as missing and "
        "impute them with the training-fold median. As Section 5 shows, this single piece of "
        "data hygiene is worth more than any model swap we tried.")
    doc.add_heading("3.2  External cohort: NHANES 2017–2018", level=2)
    add_body(doc,
        "To test whether the model travels, we needed data it had never seen, collected by "
        "different people, with different instruments, in a different era. The National Health "
        "and Nutrition Examination Survey fits. From the 2017–2018 cycle we drew the demographics, "
        "fasting-glucose, body-measurement, blood-pressure, and diabetes-questionnaire files, "
        "joined them on the respondent identifier, and restricted to women aged twenty-one and "
        "over to match the Pima inclusion criteria. After dropping records without a usable "
        "diabetes answer or the core measurements, we were left with 1,181 respondents, of whom "
        "about sixteen percent report a "
        "diabetes diagnosis. Only four features line up cleanly between the two datasets — "
        "glucose, blood pressure, BMI, and age — so the external test uses that shared subset. "
        "The base-rate gap between the two cohorts, thirty-five percent against sixteen, turns "
        "out to be the whole point.")
    doc.add_heading("3.3  Geographic context: CDC PLACES 2024", level=2)
    add_body(doc,
        "The PLACES 2024 county release gives model-based estimates of adult prevalence for "
        "every United States county. We pulled five measures that bear on diabetes risk or care: "
        "diagnosed diabetes, obesity, physical inactivity, high blood pressure, and lack of "
        "health insurance. Across the 3,144 counties, the median county reports diabetes "
        "prevalence of 13.2 percent, obesity of 38.2 percent, and physical inactivity of 27.4 "
        "percent. These are population summaries, not individual records, and we use them as "
        "context rather than as features fed to the model — a distinction we return to in "
        "Section 4.5.")

    # ---- 4 Methods ----
    doc.add_heading("4  Methods", level=1)
    doc.add_heading("4.1  Preprocessing", level=2)
    add_body(doc,
        "Every preprocessing statistic is learned inside a scikit-learn pipeline so that no "
        "information leaks from test to train. The pipeline recodes the zero-as-missing columns, "
        "imputes with the median, optionally adds a handful of derived features — products and "
        "ratios such as glucose times BMI, BMI times age, and insulin over glucose, all of them "
        "chosen because they correspond to quantities a clinician would already think about — "
        "and standardises the result. Because the imputation medians are fit on training folds "
        "only, the same object can be handed a single new patient or a county's worth of them "
        "without special casing.")
    doc.add_heading("4.2  Models and selection", level=2)
    add_body(doc,
        "We compared seven estimators: logistic regression, random forest, extremely randomised "
        "trees, gradient boosting, XGBoost, LightGBM, and a stacked ensemble that feeds the "
        "first six into a logistic meta-learner. Selection was by repeated stratified five-fold "
        "cross-validation — three repeats, fifteen fits per model — scored on ROC-AUC. We chose "
        "ROC-AUC for selection because it does not depend on a threshold, which we tune "
        "separately and later.")
    doc.add_heading("4.3  Calibration and the decision threshold", level=2)
    add_body(doc,
        "A ranking is not a probability. Tree ensembles in particular tend to be overconfident "
        "in the middle of the range, so we wrapped the chosen model in Platt scaling fit by "
        "internal cross-validation, then chose an operating threshold by maximising the "
        "F1-score on out-of-fold predictions rather than defaulting to one-half. The threshold "
        "that fell out, near 0.28, is well below the textbook cut. That is not an accident: in a "
        "screening context, missing a case is usually costlier than a false alarm, and a low "
        "threshold trades precision for the recall that screening actually cares about.")
    doc.add_heading("4.4  Evaluation protocol", level=2)
    add_body(doc,
        "We report two kinds of evidence. For model selection we use the repeated "
        "cross-validation above. For honest performance numbers we hold out twenty percent of "
        "the Pima data, never touched during selection, and attach a 95 percent confidence "
        "interval to every metric by stratified bootstrap resampling with two thousand "
        "replicates. Where we compare two models head to head we use the DeLong test, which "
        "respects the fact that both models are scored on the same patients. We consider this "
        "the least glamorous and most important part of the whole exercise.")
    doc.add_heading("4.5  Geographic prior correction", level=2)
    add_body(doc,
        "Here is the idea that ties the chapter together. The model learns to separate cases "
        "from non-cases in a cohort where roughly a third of people have diabetes. Drop that same "
        "model into a population where only one in eight does, and its raw probabilities are "
        "miscalibrated for that setting even if its ranking is perfect. Bayes' rule tells us how "
        "to fix this if we are willing to assume the ranking transports: we shift the model's "
        "odds by the ratio of the local prevalence odds to the training prevalence odds.")
    eq = doc.add_paragraph(); eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eq.add_run("odds_adjusted  =  odds_model  ×  [ p_local / (1 − p_local) ]  ÷  "
                   "[ p_train / (1 − p_train) ]")
    r.italic = True; r.font.size = Pt(11)
    add_body(doc,
        "Setting the local prevalence to a county's PLACES diabetes estimate re-expresses an "
        "individual's score against where they live. The correction is deliberately humble: it "
        "moves the base rate, not the biology, and as a sanity check, applying it at the "
        "training prevalence returns the input unchanged to within floating-point error "
        f"(maximum deviation {M['geo']['identity_map_max_error']:.0e}). The assumption it leans "
        "on — that the model's discriminative signal carries across populations — is not free, "
        "and Section 5.6 is where we go looking for evidence that it holds.")

    # ---- 5 Results ----
    doc.add_heading("5  Results", level=1)

    doc.add_heading("5.1  Which design choices mattered", level=2)
    add_body(doc,
        "We started with an ablation because we wanted to know what was doing the work. The "
        "answer was slightly humbling. Cleaning the zero-coded missing values lifted "
        "cross-validated ROC-AUC from 0.836 to 0.843; the derived clinical features nudged it to "
        "0.844; and the stacked ensemble, for all its machinery, did not improve on a single "
        "well-configured tree model at all. If you take one practical lesson from this section, "
        "it is that an hour spent understanding how the data encode missingness beat every hour "
        "we spent on model complexity.")
    add_table_from_csv(doc, "ablation_table.csv",
        "Table 1. Ablation study. Repeated stratified five-fold cross-validated ROC-AUC as each "
        "pipeline component is added.",
        col_format={"CV ROC-AUC": "{:.4f}", "Std": "{:.4f}"})

    doc.add_heading("5.2  Model comparison, and a result we did not expect to like", level=2)
    add_body(doc,
        "Extremely randomised trees came out narrowly on top, but the DeLong test had something "
        "uncomfortable to say. The gap between the best tree model and plain logistic regression "
        "is not statistically significant — the p-value sits around 0.90, which is about as "
        "clear a null result as one ever sees. The stacked ensemble is likewise "
        "indistinguishable from the winner. On data this size, the honest summary is that the "
        "simple, interpretable baseline is every bit as good as the complicated thing, and we "
        "would have been wrong to dress up the ensemble as a contribution.")
    add_table_from_csv(doc, "model_comparison.csv",
        "Table 2. Model comparison. Cross-validated ROC-AUC and the DeLong test p-value for the "
        "difference from the best model. Large p-values indicate the simpler model is "
        "statistically indistinguishable from the best.",
        col_format={"CV ROC-AUC": "{:.4f}", "Std": "{:.4f}", "DeLong p vs best": "{:.4f}"})

    doc.add_heading("5.3  Held-out performance, with the uncertainty left in", level=2)
    add_body(doc,
        "On the untouched hold-out the calibrated model reaches a ROC-AUC of 0.82, recall of "
        "0.83, and an F1 of 0.69 at the tuned threshold. The confidence intervals are wide — the "
        "ROC-AUC interval runs from roughly 0.75 to 0.88 — and they are wide for an unavoidable "
        "reason: a twenty-percent hold-out of 768 rows is about 154 patients, and you cannot wring "
        "tight intervals out of that. We would rather show the width than pretend a point estimate "
        "is more solid than it is.")
    add_table_from_csv(doc, "bootstrap_cis.csv",
        "Table 3. Held-out test metrics with 95% bootstrap confidence intervals (2,000 "
        "stratified resamples), calibrated, at the tuned decision threshold.",
        col_format={"estimate": "{:.3f}", "ci_low": "{:.3f}", "ci_high": "{:.3f}"})

    doc.add_heading("5.4  Are the probabilities trustworthy?", level=2)
    add_body(doc,
        "Calibration asks a different question from discrimination: not whether the model ranks "
        "patients correctly, but whether a stated probability of 0.3 really happens about thirty "
        "percent of the time. The reliability diagram in Figure 1 shows the raw tree-ensemble "
        "probabilities drifting away from the diagonal, and Platt scaling pulling them back "
        "toward it. The expected calibration error drops accordingly. This matters here because "
        "the geographic correction in Section 5.7 operates on probabilities, not ranks — feed it "
        "a miscalibrated score and the adjustment inherits the error.")
    add_figure(doc, "calibration.png",
        "Figure 1. Reliability diagram before and after Platt scaling. Points on the dashed "
        "diagonal are perfectly calibrated; ECE is the expected calibration error.", width=4.4)
    add_figure(doc, "roc_curves.png",
        "Figure 2. ROC curve for the held-out test set with a 95% bootstrap confidence band, "
        "conveying the same uncertainty as Table 3 in visual form.", width=4.4)

    doc.add_heading("5.5  What the model is actually looking at", level=2)
    add_body(doc,
        "We ran SHAP partly to explain the model and partly to check it had not learned "
        "something silly. It had not. Glucose dominates the attributions, followed by the "
        "glucose-times-BMI interaction, the BMI-times-age interaction, and age itself — an "
        "ordering that any clinician would have predicted before seeing a line of code. We treat "
        "that agreement as reassurance rather than discovery: when a model's reasons match "
        "established medical knowledge, a reader who does not trust machine learning has one less "
        "reason to dismiss it.")
    add_figure(doc, "shap_summary.png",
        "Figure 3. SHAP global feature importance. Each point is one patient; horizontal spread "
        "shows the feature's push toward or away from a diabetes prediction, coloured by feature "
        "value.", width=5.2)

    doc.add_heading("5.6  Does it travel? External validation on NHANES", level=2)
    iauc = ext["internal"]["roc_auc"]; eauc = ext["external"]["roc_auc"]
    eci = ext["external_ci"]["roc_auc"]
    add_body(doc,
        "This is the result that surprised us, in a good way. Trained on Pima and dropped onto "
        f"NHANES without any refitting, the model's ROC-AUC barely moves: {iauc:.2f} on the Pima "
        f"hold-out against {eauc:.2f} on NHANES, with a confidence interval of roughly "
        f"{eci['ci_low']:.2f} to {eci['ci_high']:.2f}. Discrimination, in other words, "
        "transports almost perfectly across two cohorts separated by decades, instruments, and "
        "geography. The precision-recall figures tell the complementary story: PR-AUC falls and "
        "raw accuracy actually rises on NHANES, both of which follow mechanically from the lower "
        "base rate rather than from any change in the model. We dwell on this because it is the "
        "empirical permission slip for the geographic adjustment: if the ranking holds across "
        "populations, then correcting only the base rate by place is a defensible move.")
    add_table_from_csv(doc, "external_validation.csv",
        "Table 4. Internal (Pima hold-out) versus external (NHANES) performance on the four "
        "shared features. External metrics carry 95% bootstrap confidence intervals.",
        col_format={"internal_pima": "{:.3f}", "external_nhanes": "{:.3f}",
                    "external_ci_low": "{:.3f}", "external_ci_high": "{:.3f}"})

    doc.add_heading("5.7  Risk in its place: the geographic layer", level=2)
    add_body(doc,
        "Figure 4 maps diagnosed-diabetes prevalence across every United States county, and the "
        "familiar geography of the condition is immediately visible — the band of high prevalence "
        "through the South stands out exactly as public-health data would lead you to expect. "
        "Figure 5 shows the correction at work. A raw model score of 0.6 does not mean the same "
        "thing for a patient in a county where one adult in twenty has diabetes as it does where "
        "one in five does, and the curves make that difference explicit. The dashed line is the "
        "training prevalence, where the correction does nothing by construction. We want to be "
        "precise about what this buys: it re-expresses population-level risk context, not "
        "place-based biology, and it is emphatically not a diagnosis keyed to a postcode.")
    add_figure(doc, "geo_map.png",
        "Figure 4. Adult diagnosed-diabetes prevalence by United States county (CDC PLACES "
        "2024), across all 3,144 counties.", width=5.8)
    add_figure(doc, "geo_prior_correction.png",
        "Figure 5. Prior correction across county base rates. The same model probability maps to "
        "different adjusted risks depending on local prevalence; the dashed line marks the "
        "training prevalence, where the correction is the identity.", width=4.6)

    # ---- 6 System ----
    doc.add_heading("6  The Interactive System", level=1)
    add_body(doc,
        "None of this would be much use locked in a notebook, so we wrapped it in an application "
        "built with Streamlit and a small FastAPI service for programmatic access. The dataset "
        "and a trained model ship with the app, so every page works on first launch without an "
        "upload. A user can explore the data, watch the seven models compete, drag a threshold "
        "and see the confusion matrix respond in real time, enter a single patient or upload a "
        "batch for scoring, and — the part we care about — drop that prediction onto a map of any "
        "United States county to read it against local prevalence. A separate results page "
        "exposes every table and figure in this chapter, regenerated from code, so a sceptical "
        "reader can audit the numbers rather than take them on faith. The whole analysis "
        "reproduces from a single command, with dependencies pinned and seeds fixed.")

    # ---- 7 Ethics ----
    doc.add_heading("7  Ethics, Limitations, and Responsible Use", level=1)
    add_body(doc,
        "We cannot use this dataset without naming where it comes from. The Pima, or Akimel "
        "O'odham, people of Arizona have been studied intensively for diabetes since the "
        "mid-twentieth century, and that history is not a neutral one — it raises real questions "
        "about consent, benefit, and the extraction of knowledge from a community that has not "
        "always shared in the results. A model trained on a few hundred of their records should "
        "never be turned around and used to make claims about that community, and the "
        "transportability we demonstrate on NHANES is a statement about a statistical signal, not "
        "a licence to generalise about any group of people.")
    add_body(doc,
        "The technical limitations are equally worth stating plainly. The training set is small, "
        "single-cohort, and restricted to adult women, which bounds how far any conclusion can "
        "reach. The external labels in NHANES are self-reported, and clinical definitions of "
        "diabetes have shifted across the decades separating the two datasets. The shared-feature "
        "model leans on only four variables. And the geographic correction rests on an assumption "
        "— that discrimination transports — which we tested and found support for here, but which "
        "could fail for a different condition or a more distant population. Above all, this is a "
        "screening-style illustration and an educational artifact. It is not a medical device, it "
        "has not been validated for clinical use, and no individual should make a health decision "
        "on the strength of its output.")

    # ---- 8 Conclusion ----
    doc.add_heading("8  Conclusion and Future Work", level=1)
    add_body(doc,
        "The model in this chapter is ordinary on purpose. What we hope is less ordinary is the "
        "discipline around it: an ablation that tells you what mattered, intervals that admit "
        "what we do not know, a significance test that talked us out of a result we would have "
        "preferred, an external cohort that the model had to face cold, and a geographic layer "
        "that puts an individual score in the context of a place. The standout empirical finding "
        "is that discrimination held up across a decades-wide gap between cohorts even as the "
        "base rate halved, and that finding is precisely what makes the prevalence correction "
        "more than a cosmetic touch.")
    add_body(doc,
        "Several extensions suggest themselves. PLACES publishes estimates at the census-tract "
        "level, which would sharpen the geography from county to neighbourhood. Linking in the "
        "food-access and walkability data already collected by federal agencies would let the "
        "context speak to modifiable environment rather than prevalence alone. Conformal "
        "prediction would let the system say it does not know, instead of always answering. And a "
        "larger, more representative training cohort — BRFSS runs to hundreds of thousands of "
        "records — would narrow the intervals that Section 5 is so insistent about. The thread "
        "running through all of these is the same one running through the chapter: a prediction "
        "is more honest, and more useful, when it carries its uncertainty and its context along "
        "with it.")

    # ---- References ----
    doc.add_heading("References", level=1)
    refs = [
        "Smith, J. W., Everhart, J. E., Dickson, W. C., Knowler, W. C., & Johannes, R. S. "
        "(1988). Using the ADAP learning algorithm to forecast the onset of diabetes mellitus. "
        "Proceedings of the Annual Symposium on Computer Application in Medical Care, 261–265.",
        "Geurts, P., Ernst, D., & Wehenkel, L. (2006). Extremely randomized trees. Machine "
        "Learning, 63(1), 3–42.",
        "Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. Journal of "
        "Machine Learning Research, 12, 2825–2830.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings "
        "of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, "
        "785–794.",
        "Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. "
        "Advances in Neural Information Processing Systems, 30.",
        "Platt, J. (1999). Probabilistic outputs for support vector machines and comparisons to "
        "regularized likelihood methods. Advances in Large Margin Classifiers, 10(3), 61–74.",
        "DeLong, E. R., DeLong, D. M., & Clarke-Pearson, D. L. (1988). Comparing the areas under "
        "two or more correlated receiver operating characteristic curves: a nonparametric "
        "approach. Biometrics, 44(3), 837–845.",
        "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model "
        "predictions. Advances in Neural Information Processing Systems, 30, 4765–4774.",
        "Vickers, A. J., & Elkin, E. B. (2006). Decision curve analysis: a novel method for "
        "evaluating prediction models. Medical Decision Making, 26(6), 565–574.",
        "Centers for Disease Control and Prevention. (2024). PLACES: Local Data for Better "
        "Health, County Data. U.S. Department of Health and Human Services.",
        "Centers for Disease Control and Prevention, National Center for Health Statistics. "
        "(2020). National Health and Nutrition Examination Survey Data, 2017–2018 cycle.",
        "Efron, B., & Tibshirani, R. J. (1993). An Introduction to the Bootstrap. Chapman & "
        "Hall/CRC.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        r = p.add_run(f"[{i}]  {ref}")
        r.font.size = Pt(9.5)

    doc.save(OUT)
    print(f"Saved chapter -> {OUT}")
    print(f"Words (approx): {sum(len(p.text.split()) for p in doc.paragraphs)}")


if __name__ == "__main__":
    build()
